"""Backend-only services for the Examination paper bank."""
from __future__ import annotations

import base64
import difflib
import hashlib
import json
import os
import random
import re
import time
import unicodedata
from collections import Counter
from io import BytesIO
from pathlib import Path

import requests
from cryptography.fernet import Fernet, InvalidToken
from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from jsonschema import Draft202012Validator
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader

from .models import AiProviderConfig, AiUsageLog, BlueprintSlot, ExamPaper, ExamQuestion, ExamReview, ExamSourceDocument

DIFFICULTIES = ('EASY', 'MEDIUM', 'HARD', 'VERY_HARD')
DIFFICULTY_LABELS = {'EASY': 'Dễ', 'MEDIUM': 'Trung bình', 'HARD': 'Khó', 'VERY_HARD': 'Rất khó'}
CHECK_STATUSES = {'PENDING', 'PASSED', 'AI_FIXED', 'NEEDS_REVIEW', 'WARNING'}
QUESTION_SCHEMA = {
    'type': 'object', 'required': ['content', 'choices', 'correctAnswer', 'difficulty', 'topic'],
    'properties': {
        'content': {'type': 'string', 'minLength': 3},
        'choices': {'type': 'array', 'minItems': 4, 'maxItems': 4, 'items': {'type': 'string', 'minLength': 1}},
        'correctAnswer': {'type': 'string', 'enum': ['A', 'B', 'C', 'D']},
        'explanation': {'type': 'string'},
        'difficulty': {'type': 'string', 'enum': list(DIFFICULTIES)},
        'topic': {'type': 'string'},
    },
}
QUESTION_VALIDATOR = Draft202012Validator(QUESTION_SCHEMA)


def _cipher() -> Fernet:
    key = os.getenv('AI_CONFIG_ENCRYPTION_KEY', '').strip()
    if not key:
        if not settings.DEBUG:
            raise ValueError('Thiếu AI_CONFIG_ENCRYPTION_KEY trên máy chủ.')
        key = base64.urlsafe_b64encode(hashlib.sha256(settings.SECRET_KEY.encode()).digest()).decode()
    return Fernet(key.encode())


def encrypt_secret(value: str) -> str:
    return _cipher().encrypt(value.encode()).decode() if value else ''


def decrypt_secret(value: str) -> str:
    if not value:
        return ''
    try:
        return _cipher().decrypt(value.encode()).decode()
    except (InvalidToken, ValueError) as exc:
        raise ValueError('Không thể giải mã API key AI.') from exc


def masked_secret(value: str) -> str:
    if not value:
        return ''
    plain = decrypt_secret(value)
    return f'{plain[:4]}••••••••{plain[-4:]}' if len(plain) > 8 else '••••••••'


def default_distribution(total: int) -> dict[str, int]:
    total = max(0, int(total or 0))
    easy = total * 20 // 100
    hard = total * 20 // 100
    very_hard = total * 10 // 100
    return {'EASY': easy, 'MEDIUM': total - easy - hard - very_hard, 'HARD': hard, 'VERY_HARD': very_hard}


def validate_distribution(total: int, values: dict) -> dict[str, int]:
    normalized = {key: max(0, int((values or {}).get(key, 0) or 0)) for key in DIFFICULTIES}
    if sum(normalized.values()) != int(total or 0):
        raise ValueError('Tổng câu Dễ, Trung bình, Khó và Rất khó phải bằng tổng số câu.')
    return normalized


def normalize_question(value: dict, order: int = 1) -> dict:
    raw = dict(value or {})
    choices = raw.get('choices') or raw.get('options') or []
    if isinstance(choices, dict):
        choices = [choices.get(label, '') for label in 'ABCD']
    if isinstance(choices, list):
        choices = [str(item.get('text', '')) if isinstance(item, dict) else str(item) for item in choices]
    normalized = {
        'content': str(raw.get('content') or raw.get('question') or '').strip(),
        'choices': choices,
        'correctAnswer': str(raw.get('correctAnswer') or raw.get('correct_answer') or '').strip().upper(),
        'explanation': str(raw.get('explanation') or '').strip(),
        'difficulty': str(raw.get('difficulty') or 'MEDIUM').strip().upper(),
        'topic': str(raw.get('topic') or '').strip(),
        'order': int(raw.get('order') or order),
    }
    errors = sorted(QUESTION_VALIDATOR.iter_errors(normalized), key=lambda error: list(error.path))
    if errors:
        raise ValueError('Câu hỏi không hợp lệ: ' + '; '.join(error.message for error in errors[:3]))
    return normalized


def serialize_question(question: ExamQuestion) -> dict:
    return {
        'id': str(question.id), 'order': question.order, 'content': question.content,
        'choices': question.choices or [], 'correctAnswer': question.correct_answer,
        'explanation': question.explanation, 'difficulty': question.difficulty, 'topic': question.topic,
        'checkStatus': question.check_status, 'warnings': question.warnings or [], 'aiMetadata': question.ai_metadata or {},
        'blueprintSlotId': str(question.blueprint_slot_id or ''), 'questionType': question.question_type, 'score': float(question.score), 'slotMetadata': question.slot_metadata or {},
    }


def serialize_source(source: ExamSourceDocument) -> dict:
    return {
        'id': str(source.id), 'type': source.source_type, 'name': source.name,
        'fileUrl': source.file.url if source.file else '', 'referencedPaperId': str(source.referenced_paper_id or ''),
        'createdAt': source.created_at.isoformat(),
    }


def serialize_paper(paper: ExamPaper, include_questions: bool = False) -> dict:
    latest_job = paper.generation_jobs.order_by('-created_at').first()
    data = {
        'id': str(paper.id), 'title': paper.title, 'blueprintVersionId': str(paper.blueprint_version_id or ''), 'blueprintName': paper.blueprint_version.blueprint.name if paper.blueprint_version_id else 'Đề cũ chưa có ma trận', 'blueprintVersion': paper.blueprint_version.version_number if paper.blueprint_version_id else None,
        'competitionId': paper.competition_id or '', 'competitionName': paper.competition.name if paper.competition else '',
        'sessionId': paper.session_id or '', 'sessionName': paper.session.name if paper.session else '',
        'subject': paper.subject, 'gradeOrCategory': paper.grade_or_category, 'language': paper.language,
        'durationMinutes': paper.duration_minutes, 'totalQuestions': paper.total_questions,
        'difficultyDistribution': paper.difficulty_distribution or default_distribution(paper.total_questions),
        'status': paper.status, 'description': paper.description, 'aiGenerationStatus': paper.ai_generation_status,
        'aiGenerationMessage': paper.ai_generation_message, 'createdBy': paper.created_by, 'updatedBy': paper.updated_by,
        'qualityReport': paper.quality_report or {}, 'workflowLog': paper.workflow_log or [],
        'aiChatHistory': paper.ai_chat_history or [],
        'generationProgress': {
            'generated': latest_job.generated_count if latest_job else 0,
            'total': paper.total_questions,
            'resumable': bool(latest_job and latest_job.status == 'FAILED' and latest_job.partial_questions),
        },
        'draftExportedAt': paper.draft_exported_at.isoformat() if paper.draft_exported_at else None,
        'approvedAt': paper.approved_at.isoformat() if paper.approved_at else None, 'approvedBy': paper.approved_by,
        'officialExportedAt': paper.official_exported_at.isoformat() if paper.official_exported_at else None,
        'bankedAt': paper.banked_at.isoformat() if paper.banked_at else None,
        'createdAt': paper.created_at.isoformat(), 'updatedAt': paper.updated_at.isoformat(),
        'questionCount': paper.questions.count(),
    }
    if include_questions:
        data['questions'] = [serialize_question(item) for item in paper.questions.all()]
        data['sources'] = [serialize_source(item) for item in paper.sources.select_related('referenced_paper').all()]
        data['reviews'] = [{
            'id': str(item.id), 'scope': item.scope, 'questionId': str(item.question_id or ''),
            'verdict': item.verdict, 'notes': item.notes, 'checks': item.checks or {},
            'reviewer': item.reviewer, 'updatedAt': item.updated_at.isoformat(),
        } for item in paper.human_reviews.select_related('question').all()]
    return data


def _word_math_text(element) -> str:
    """Convert common Office Math nodes to compact, prompt-friendly plain text."""
    tag = element.tag.rsplit('}', 1)[-1]
    children = list(element)
    if tag in {'t', 'instrText'}:
        return element.text or ''
    if tag == 'tab':
        return '\t'
    if tag in {'br', 'cr'}:
        return '\n'
    by_tag = {child.tag.rsplit('}', 1)[-1]: child for child in children}
    if tag == 'f' and 'num' in by_tag and 'den' in by_tag:
        return f"({_word_math_text(by_tag['num'])})/({_word_math_text(by_tag['den'])})"
    if tag == 'sSup' and 'e' in by_tag and 'sup' in by_tag:
        return f"{_word_math_text(by_tag['e'])}^({_word_math_text(by_tag['sup'])})"
    if tag == 'sSub' and 'e' in by_tag and 'sub' in by_tag:
        return f"{_word_math_text(by_tag['e'])}_({_word_math_text(by_tag['sub'])})"
    if tag == 'rad' and 'e' in by_tag:
        return f"√({_word_math_text(by_tag['e'])})"
    return ''.join(_word_math_text(child) for child in children)


def _paragraph_source_text(paragraph: Paragraph) -> str:
    return _word_math_text(paragraph._p).strip()


def _docx_source_lines(document: Document) -> list[str]:
    """Read paragraphs and tables in document order instead of dropping tables."""
    lines = []
    for child in document.element.body.iterchildren():
        if child.tag == qn('w:p'):
            text = _paragraph_source_text(Paragraph(child, document))
            if text:
                lines.append(text)
        elif child.tag == qn('w:tbl'):
            table = Table(child, document)
            for row in table.rows:
                values = []
                for cell in row.cells:
                    value = ' '.join(filter(None, (_paragraph_source_text(item) for item in cell.paragraphs))).strip()
                    if not values or value != values[-1]:
                        values.append(value)
                if any(values):
                    lines.append(' | '.join(values))
    return lines


def read_uploaded_source(uploaded) -> str:
    extension = Path(uploaded.name).suffix.lower()
    if extension not in {'.pdf', '.docx', '.xlsx'}:
        raise ValueError('Chỉ hỗ trợ tệp PDF, DOCX hoặc XLSX.')
    if uploaded.size > 15 * 1024 * 1024:
        raise ValueError('Mỗi tệp nguồn tối đa 15 MB.')
    content = uploaded.read()
    uploaded.seek(0)
    if extension == '.pdf':
        return '\n'.join((page.extract_text() or '') for page in PdfReader(BytesIO(content)).pages)[:120000]
    if extension == '.docx':
        return '\n'.join(_docx_source_lines(Document(BytesIO(content))))[:120000]
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    rows = []
    for sheet in workbook.worksheets:
        rows.append(f'[{sheet.title}]')
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value not in (None, '')]
            if values:
                rows.append(' | '.join(values))
    return '\n'.join(rows)[:120000]


def source_context(paper: ExamPaper, max_chars: int = 30000) -> str:
    """Build a bounded, balanced context so one long source cannot consume every AI request."""
    max_chars = max(1000, min(int(max_chars), 60000))
    chunks = [f'Yêu cầu bổ sung: {paper.description}'] if paper.description else []
    sources = list(paper.sources.select_related('referenced_paper').all())
    reserved = sum(len(chunk) + 2 for chunk in chunks)
    per_source = max(1000, (max_chars - reserved) // max(1, len(sources)))
    for source in sources:
        if source.source_type == 'PAPER' and source.referenced_paper:
            questions = '\n'.join(f'- {question.content}' for question in source.referenced_paper.questions.all()[:100])
            chunk = f'Đề tham chiếu: {source.referenced_paper.title}\n{questions}'
        elif source.extracted_text:
            chunk = f'Nguồn {source.name}:\n{source.extracted_text}'
        else:
            continue
        chunks.append(chunk[:per_source])
    return '\n\n'.join(chunks)[:max_chars]


def _config() -> AiProviderConfig:
    return AiProviderConfig.objects.filter(provider='openai', is_enabled=True).order_by('priority', 'id').first() or AiProviderConfig(provider='openai')


def _active_configs() -> list[AiProviderConfig]:
    return list(AiProviderConfig.objects.filter(provider='openai', is_enabled=True).exclude(health_status=AiProviderConfig.STATUS_EXHAUSTED).order_by('priority', 'id'))


def serialize_ai_config(config: AiProviderConfig) -> dict:
    return {
        'id': config.pk, 'name': config.name, 'provider': config.provider, 'baseUrl': config.base_url, 'apiKeyMasked': masked_secret(config.api_key_encrypted),
        'hasApiKey': bool(config.api_key_encrypted), 'generationModel': config.generation_model,
        'reviewModel': config.review_model, 'temperature': config.temperature, 'maxTokens': config.max_tokens,
        'timeoutSeconds': config.timeout_seconds, 'maxRetries': config.max_retries,
        'priority': config.priority, 'isEnabled': config.is_enabled, 'healthStatus': config.health_status,
        'lastError': config.last_error, 'lastFailedAt': config.last_failed_at.isoformat() if config.last_failed_at else None,
        'lastUsedAt': config.last_used_at.isoformat() if config.last_used_at else None,
        'updatedAt': config.updated_at.isoformat() if config.pk else '',
    }


def call_ai_json(*, paper: ExamPaper, task_type: str, model: str, system: str, prompt: str, user_email: str, max_tokens: int | None = None) -> dict:
    configs = _active_configs()
    if not configs:
        if AiProviderConfig.objects.filter(provider='openai', is_enabled=True, health_status=AiProviderConfig.STATUS_EXHAUSTED).exists():
            raise ValueError('Tất cả cấu hình AI đang bật đều đã hết hạn mức. Tiến trình đã được lưu nháp; hãy bổ sung hạn mức hoặc cập nhật API key rồi tiếp tục.')
        raise ValueError('Chưa có cấu hình AI đang hoạt động. Vui lòng vào Cấu hình AI để thêm API key.')
    failures: list[str] = []
    review_task = task_type in {'review', 'review_auto_fix'}
    quota_codes = {'credit_balance_exhausted', 'organization_spend_limit_exceeded', 'project_spend_limit_exceeded', 'organization_usage_limit_exceeded'}
    for config in configs:
        started = time.monotonic()
        selected_model = (config.review_model if review_task else config.generation_model) or model
        try:
            api_key = decrypt_secret(config.api_key_encrypted)
        except ValueError as exc:
            api_key = ''
            failures.append(f'{config.name}: {exc}')
        if not api_key:
            config.health_status = AiProviderConfig.STATUS_ERROR
            config.last_error = 'Chưa có API key hợp lệ.'
            config.last_failed_at = timezone.now()
            config.save(update_fields=['health_status', 'last_error', 'last_failed_at', 'updated_at'])
            failures.append(f'{config.name}: chưa có API key hợp lệ')
            continue
        endpoint = config.base_url.rstrip('/') + '/chat/completions'
        payload = {
            'model': selected_model, 'temperature': config.temperature,
            'max_tokens': min(config.max_tokens, max_tokens) if max_tokens else config.max_tokens,
            'response_format': {'type': 'json_object'},
            'messages': [{'role': 'system', 'content': system}, {'role': 'user', 'content': prompt}],
        }
        last_error = ''
        should_failover = False
        for attempt in range(config.max_retries + 1):
            try:
                response = requests.post(endpoint, headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'}, json=payload, timeout=config.timeout_seconds)
                if response.ok:
                    body = response.json()
                    content = body['choices'][0]['message']['content']
                    parsed = json.loads(content)
                    usage = body.get('usage') or {}
                    config.health_status = AiProviderConfig.STATUS_READY
                    config.last_error = ''
                    config.last_used_at = timezone.now()
                    config.save(update_fields=['health_status', 'last_error', 'last_used_at', 'updated_at'])
                    AiUsageLog.objects.create(config=config, paper=paper, user_email=user_email, task_type=task_type, provider=config.provider, model=selected_model, input_tokens=int(usage.get('prompt_tokens') or 0), output_tokens=int(usage.get('completion_tokens') or 0), duration_ms=int((time.monotonic() - started) * 1000), success=True)
                    if failures:
                        paper._ai_route_notice = f'Đã tự động chuyển sang {config.name}. ' + ' · '.join(failures)
                    return parsed
                try:
                    error_body = response.json().get('error') or {}
                except Exception:
                    error_body = {}
                error_code = str(error_body.get('code') or '')
                error_type = str(error_body.get('type') or '')
                error_message = str(error_body.get('message') or response.text or f'HTTP {response.status_code}')[:1000]
                last_error = f'{response.status_code} {error_code or error_type}: {error_message}'.strip()
                exhausted = response.status_code == 429 and (error_code in quota_codes or error_type == 'insufficient_quota')
                retryable = response.status_code == 429 and not exhausted or response.status_code >= 500
                should_failover = exhausted or retryable or response.status_code in {401, 403}
                if exhausted:
                    config.health_status = AiProviderConfig.STATUS_EXHAUSTED
                    break
                if retryable and attempt < config.max_retries:
                    time.sleep(min(8, 2 ** attempt) + random.random())
                    continue
                if should_failover:
                    config.health_status = AiProviderConfig.STATUS_ERROR
                    break
                response.raise_for_status()
            except (requests.Timeout, requests.ConnectionError) as exc:
                last_error = str(exc)
                should_failover = True
                if attempt < config.max_retries:
                    time.sleep(min(8, 2 ** attempt) + random.random())
                    continue
                config.health_status = AiProviderConfig.STATUS_ERROR
                break
            except Exception as exc:
                last_error = str(exc)
                should_failover = False
                config.health_status = AiProviderConfig.STATUS_ERROR
                break
        config.last_error = last_error[:4000]
        config.last_failed_at = timezone.now()
        config.save(update_fields=['health_status', 'last_error', 'last_failed_at', 'updated_at'])
        AiUsageLog.objects.create(config=config, paper=paper, user_email=user_email, task_type=task_type, provider=config.provider, model=selected_model, duration_ms=int((time.monotonic() - started) * 1000), success=False, error_message=last_error[:4000])
        failures.append(f'{config.name}: {"hết hạn mức" if config.health_status == AiProviderConfig.STATUS_EXHAUSTED else last_error}')
        if not should_failover:
            break
    paper._ai_route_notice = ' · '.join(failures)
    raise ValueError('AI chưa thể hoàn tất yêu cầu. Tiến trình đã được lưu nháp. ' + ' · '.join(failures))


def generate_questions_with_ai(paper: ExamPaper, user_email: str) -> list[dict]:
    distribution = validate_distribution(paper.total_questions, paper.difficulty_distribution)
    demand = ', '.join(f'{DIFFICULTY_LABELS[key]}: {value}' for key, value in distribution.items())
    prompt = f'''Tạo chính xác {paper.total_questions} câu trắc nghiệm một đáp án đúng cho đề "{paper.title}".\nCuộc thi: {paper.competition.name if paper.competition else ''}\nKỳ thi: {paper.session.name if paper.session else ''}\nMôn: {paper.subject}; Khối/Bảng: {paper.grade_or_category}; Ngôn ngữ: {paper.language}; Thời gian: {paper.duration_minutes} phút.\nPhân bố: {demand}.\nChỉ trả JSON theo dạng {{"questions":[{{"content":"","choices":["A","B","C","D"],"correctAnswer":"A","explanation":"","difficulty":"EASY|MEDIUM|HARD|VERY_HARD","topic":""}}]}}. Không sao chép nguyên văn nguồn.\n\nNguồn tham chiếu:\n{source_context(paper)}'''
    result = call_ai_json(paper=paper, task_type='generate', model=_config().generation_model, user_email=user_email, system='Bạn là chuyên gia ra đề. Chỉ trả JSON hợp lệ, không markdown, không văn bản tự do.', prompt=prompt)
    rows = result.get('questions') if isinstance(result, dict) else None
    if not isinstance(rows, list) or len(rows) != paper.total_questions:
        raise ValueError('AI trả về số lượng câu hỏi không đúng yêu cầu.')
    normalized = [normalize_question(item, index + 1) for index, item in enumerate(rows)]
    actual = {key: sum(1 for item in normalized if item['difficulty'] == key) for key in DIFFICULTIES}
    if actual != distribution:
        raise ValueError('AI trả về phân bố độ khó không đúng yêu cầu.')
    return normalized


def review_questions_with_ai(paper: ExamPaper, auto_fix: bool, user_email: str) -> list[dict]:
    current = [serialize_question(item) for item in paper.questions.all()]
    prompt = f'''Kiểm tra đề dưới đây: đúng loại câu và đúng số phương án theo từng slot, một đáp án đúng, số lượng, độ khó, câu trùng, mơ hồ, vượt phạm vi, lỗi ngôn ngữ và phương án nhiễu. Nếu autoFix=true, chỉ viết lại các câu lỗi và giữ nguyên mọi ràng buộc slot.\nTrả JSON {{"questions":[{{"id":"","status":"PASSED|AI_FIXED|NEEDS_REVIEW|WARNING","warnings":[""],"replacement":null hoặc question schema}}]}}.\nautoFix={str(auto_fix).lower()}\nĐề: {json.dumps(current, ensure_ascii=False)}\nNguồn: {source_context(paper)}'''
    result = call_ai_json(paper=paper, task_type='review_auto_fix' if auto_fix else 'review', model=_config().review_model, user_email=user_email, system='Bạn là kiểm định viên đề thi. Chỉ trả JSON hợp lệ.', prompt=prompt)
    return result.get('questions') if isinstance(result, dict) and isinstance(result.get('questions'), list) else []


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn('w:shd')) or OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    if shading.getparent() is None:
        properties.append(shading)


def _set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Trang ')
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText'); instruction.set(qn('xml:space'), 'preserve'); instruction.text = ' PAGE '
    separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate')
    value = OxmlElement('w:t'); value.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for item in (begin, instruction, separate, value, end):
        run._r.append(item)


def _configure_exam_document(document, paper: ExamPaper) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = Mm(20)
    section.left_margin, section.right_margin = Mm(25), Mm(20)
    section.header_distance = section.footer_distance = Mm(12.7)
    normal = document.styles['Normal']
    normal.font.name, normal.font.size = 'Times New Roman', Pt(14)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05
    document.core_properties.title = paper.title
    document.core_properties.subject = f'{paper.subject} - {paper.grade_or_category}'.strip(' -')
    document.core_properties.author = paper.created_by or 'FermatTech'
    _set_page_number(section.footer.paragraphs[0])


def _add_exam_heading(document, paper: ExamPaper, *, answer_key: bool = False) -> None:
    blueprint = paper.blueprint_version.blueprint if paper.blueprint_version_id else None
    round_name = blueprint.round_name if blueprint else ''
    competition = paper.competition.name if paper.competition else 'CUỘC THI'
    title = 'ĐÁP ÁN' if answer_key else 'ĐỀ THI'
    for text, size in ((competition.upper(), 13), (f'{title} {round_name}'.strip().upper(), 16), (f'{paper.subject} · {paper.grade_or_category}'.strip(' ·'), 14)):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold, run.font.name, run.font.size = True, 'Times New Roman', Pt(size)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    meta.add_run(f'Thời gian làm bài: {paper.duration_minutes} phút').italic = True
    if not answer_key:
        info = document.add_table(rows=2, cols=2)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER
        info.autofit = True
        values = ('Họ và tên: ........................................................', 'Số báo danh: ........................', 'Trường: ..............................................................', 'Lớp: ........................')
        for cell, value in zip((cell for row in info.rows for cell in row.cells), values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        instruction = document.add_paragraph()
        instruction.paragraph_format.space_before, instruction.paragraph_format.space_after = Pt(6), Pt(6)
        instruction.add_run('Hướng dẫn: ').bold = True
        instruction.add_run('Chọn một đáp án đúng cho mỗi câu trắc nghiệm; ghi đáp số vào ô trống đối với câu điền đáp số.')


def _add_exam_questions(document, paper: ExamPaper) -> None:
    for question in paper.questions.all():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.keep_with_next = bool(question.choices)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.add_run(f'Câu {question.order}. ').bold = True
        paragraph.add_run(question.content)
        if question.question_type == 'numeric_input':
            answer = document.add_paragraph('Đáp số:  ........................................................')
            answer.paragraph_format.left_indent = Mm(8)
            answer.paragraph_format.space_after = Pt(5)
            continue
        for index, choice in enumerate(question.choices or []):
            option = document.add_paragraph()
            option.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            option.paragraph_format.left_indent = Mm(8)
            option.paragraph_format.first_line_indent = Mm(-6)
            option.paragraph_format.space_after = Pt(1)
            option.add_run(f'{chr(65 + index)}. ').bold = True
            option.add_run(str(choice))


def _add_answer_key(document, paper: ExamPaper) -> None:
    _add_exam_heading(document, paper, answer_key=True)
    questions = list(paper.questions.all())
    columns = min(8, max(1, len(questions)))
    table = document.add_table(rows=0, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for start in range(0, len(questions), columns):
        group = questions[start:start + columns]
        number_row, answer_row = table.add_row(), table.add_row()
        for index in range(columns):
            number_cell, answer_cell = number_row.cells[index], answer_row.cells[index]
            if index < len(group):
                question = group[index]
                number_cell.text, answer_cell.text = f'Câu {question.order}', question.correct_answer
            else:
                number_cell.text = answer_cell.text = ''
            number_cell.vertical_alignment = answer_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_shading(number_cell, 'DCE6F1')
            for cell in (number_cell, answer_cell):
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
            for run in number_cell.paragraphs[0].runs + answer_cell.paragraphs[0].runs:
                run.bold = True
    explanations = [question for question in questions if question.explanation.strip()]
    if explanations:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.add_run('HƯỚNG DẪN GIẢI').bold = True
        for question in explanations:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.add_run(f'Câu {question.order}. ').bold = True
            paragraph.add_run(question.explanation)


def document_export(paper: ExamPaper, mode: str, publication: str = 'draft') -> bytes:
    """Build the previewed paper deterministically; this path never calls AI."""
    document = Document()
    _configure_exam_document(document, paper)
    if publication == 'draft':
        banner = document.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner.paragraph_format.space_after = Pt(8)
        run = banner.add_run('BẢN NHÁP — KHÔNG PHÁT HÀNH')
        run.bold, run.font.name, run.font.size = True, 'Times New Roman', Pt(14)
        run.font.color.rgb = RGBColor(192, 0, 0)
    if mode != 'answers':
        _add_exam_heading(document, paper)
        _add_exam_questions(document, paper)
    if mode in {'answers', 'combined'}:
        if mode == 'combined':
            document.add_page_break()
        _add_answer_key(document, paper)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def xlsx_export(paper: ExamPaper) -> bytes:
    workbook = Workbook(); sheet = workbook.active; sheet.title = 'Đề thi'
    sheet.append(['STT', 'Nội dung câu hỏi', 'Phương án A', 'Phương án B', 'Phương án C', 'Phương án D', 'Phương án E', 'Đáp án', 'Lời giải', 'Mức độ khó', 'Chủ đề', 'Cuộc thi', 'Trạng thái', 'Cảnh báo'])
    for question in paper.questions.all():
        choices = list(question.choices or []) + ['', '', '', '', '']
        sheet.append([question.order, question.content, choices[0], choices[1], choices[2], choices[3], choices[4], question.correct_answer, question.explanation, DIFFICULTY_LABELS.get(question.difficulty, question.difficulty), question.topic, paper.competition.name if paper.competition else '', question.check_status, '; '.join(question.warnings or [])])
    output = BytesIO(); workbook.save(output); return output.getvalue()

def revise_question_with_ai(paper: ExamPaper, question: ExamQuestion, action: str, user_email: str) -> dict:
    instructions = {
        'rewrite': 'Viết lại câu hỏi để rõ ràng hơn nhưng giữ mục tiêu kiến thức và độ khó.',
        'distractors': 'Giữ nội dung và đáp án đúng, thay ba phương án nhiễu hợp lý, có sức phân loại.',
        'increase-difficulty': 'Tăng độ khó đúng một bậc nếu có thể, không vượt phạm vi đề.',
        'decrease-difficulty': 'Giảm độ khó đúng một bậc nhưng vẫn kiểm tra cùng kiến thức.',
        'replace': 'Thay bằng một câu hoàn toàn khác, cùng chủ đề và độ khó.',
    }
    if action not in instructions:
        raise ValueError('Thao tác AI cho câu hỏi không hợp lệ.')
    prompt = f'''{instructions[action]}\nChỉ trả JSON {{"question":{{"content":"","choices":["","","",""],"correctAnswer":"A","explanation":"","difficulty":"EASY|MEDIUM|HARD|VERY_HARD","topic":""}}}}.\nCâu hiện tại: {json.dumps(serialize_question(question), ensure_ascii=False)}\nNguồn: {source_context(paper)}'''
    result = call_ai_json(paper=paper, task_type=f'question_{action}', model=_config().generation_model, user_email=user_email, system='Bạn là chuyên gia ra đề. Chỉ trả JSON hợp lệ, không markdown.', prompt=prompt)
    row = result.get('question') if isinstance(result, dict) else None
    if not isinstance(row, dict):
        raise ValueError('AI không trả về câu hỏi hợp lệ.')
    return normalize_question({**row, 'order': question.order}, question.order)
# Blueprint-aware generation overrides. Every AI call receives one immutable slot;
# the returned question is normalised back against that slot rather than AI choices.
def _text_fingerprint(value: str, *, remove_numbers: bool = False) -> str:
    text = unicodedata.normalize('NFKC', str(value or '')).lower()
    if remove_numbers:
        text = re.sub(r'\d+(?:[.,]\d+)?', '#', text)
    text = re.sub(r'[^\w#]+', ' ', text, flags=re.UNICODE)
    return ' '.join(text.split())


def _duplicate_pair(left: dict, right: dict) -> tuple[bool, str]:
    exact_left = _text_fingerprint(left.get('content', ''))
    exact_right = _text_fingerprint(right.get('content', ''))
    if exact_left and exact_left == exact_right:
        return True, 'trùng nội dung'
    skeleton_left = _text_fingerprint(left.get('content', ''), remove_numbers=True)
    skeleton_right = _text_fingerprint(right.get('content', ''), remove_numbers=True)
    if skeleton_left and skeleton_left == skeleton_right:
        return True, 'chỉ thay số'
    similarity = difflib.SequenceMatcher(None, skeleton_left, skeleton_right).ratio()
    if min(len(skeleton_left), len(skeleton_right)) >= 45 and similarity >= 0.9:
        return True, f'quá giống nhau ({similarity:.0%})'
    return False, ''


def _balanced_answer_targets(slots: list[BlueprintSlot]) -> dict[int, str]:
    counts: dict[int, Counter] = {}
    targets: dict[int, str] = {}
    previous = ''
    for sequence, slot in enumerate(slots):
        if slot.question_type != 'single_choice':
            continue
        size = max(2, int(slot.option_count or 4))
        labels = [chr(65 + index) for index in range(size)]
        counter = counts.setdefault(size, Counter())
        offset = (sequence * 2) % size
        preference = labels[offset:] + labels[:offset]
        target = min(preference, key=lambda label: (counter[label], label == previous, preference.index(label)))
        targets[slot.position] = target
        counter[target] += 1
        previous = target
    return targets


def _place_correct_answer(row: dict, target: str) -> dict:
    choices = list(row.get('choices') or [])
    current = str(row.get('correctAnswer') or '').upper()
    valid = [chr(65 + index) for index in range(len(choices))]
    if current not in valid or target not in valid:
        raise ValueError('AI chưa xác định đáp án đúng hợp lệ.')
    old_index, new_index = ord(current) - 65, ord(target) - 65
    choices[old_index], choices[new_index] = choices[new_index], choices[old_index]
    return {**row, 'choices': choices, 'correctAnswer': target}


def paper_quality_report(paper: ExamPaper, rows: list[dict] | None = None) -> dict:
    values = rows if rows is not None else [serialize_question(item) for item in paper.questions.all()]
    issues, duplicates = [], []
    for index, row in enumerate(values):
        choices = [str(item).strip() for item in (row.get('choices') or [])]
        if row.get('questionType', 'single_choice') == 'single_choice':
            if len(set(_text_fingerprint(item) for item in choices)) != len(choices):
                issues.append(f'Câu {index + 1} có phương án trùng nhau.')
            valid = [chr(65 + item) for item in range(len(choices))]
            if str(row.get('correctAnswer') or '') not in valid:
                issues.append(f'Câu {index + 1} có đáp án không hợp lệ.')
        for previous in range(index):
            duplicate, reason = _duplicate_pair(values[previous], row)
            if duplicate:
                duplicates.append({'questions': [previous + 1, index + 1], 'reason': reason})
    answers = [str(row.get('correctAnswer') or '') for row in values if row.get('questionType', 'single_choice') == 'single_choice']
    distribution = dict(sorted(Counter(answers).items()))
    max_run, run, previous_answer = 0, 0, None
    for answer in answers:
        run = run + 1 if answer == previous_answer else 1
        max_run, previous_answer = max(max_run, run), answer
    matrix_issues = []
    slots = list(paper.blueprint_version.slots.all().order_by('position')) if paper.blueprint_version_id else []
    if slots and len(values) != len(slots):
        matrix_issues.append(f'Số câu {len(values)} không khớp {len(slots)} slot ma trận.')
    for index, (row, slot) in enumerate(zip(values, slots)):
        if row.get('difficulty') != slot.difficulty or row.get('questionType', 'single_choice') != slot.question_type:
            matrix_issues.append(f'Câu {index + 1} không khớp loại câu/độ khó của slot.')
    if duplicates:
        issues.append(f'Có {len(duplicates)} cặp câu trùng hoặc chỉ thay số.')
    if max_run > 3:
        issues.append(f'Có chuỗi {max_run} đáp án liên tiếp giống nhau.')
    if distribution and max(distribution.values()) - min(distribution.values()) > 1:
        issues.append('Đáp án trắc nghiệm chưa được phân bố đều.')
    issues.extend(matrix_issues)
    return {
        'passed': not issues, 'issues': issues, 'duplicatePairs': duplicates,
        'answerDistribution': distribution, 'maxAnswerRun': max_run,
        'questionCount': len(values), 'matrixIssues': matrix_issues,
    }


def normalize_question_for_slot(value: dict, order: int, slot: BlueprintSlot) -> dict:
    raw = dict(value or {})
    question_type = slot.question_type
    choices = raw.get('choices') or raw.get('options') or []
    if isinstance(choices, dict):
        choices = [choices.get(chr(65 + index), '') for index in range(slot.option_count)]
    if isinstance(choices, list):
        choices = [str(item.get('text', '')) if isinstance(item, dict) else str(item) for item in choices]
    answer = str(raw.get('correctAnswer') or raw.get('correct_answer') or '').strip().upper()
    if question_type == 'single_choice':
        if len(choices) != slot.option_count or any(not choice.strip() for choice in choices):
            raise ValueError(f'Slot {slot.position} cần đúng {slot.option_count} phương án.')
        valid = [chr(65 + index) for index in range(slot.option_count)]
        if answer not in valid:
            raise ValueError(f'Slot {slot.position} có đáp án không hợp lệ.')
    else:
        choices = []
        if not answer:
            raise ValueError(f'Slot {slot.position} cần đáp số đúng.')
    content = str(raw.get('content') or raw.get('question') or '').strip()
    if len(content) < 3:
        raise ValueError(f'Slot {slot.position} chưa có nội dung hợp lệ.')
    return {'content': content, 'choices': choices, 'correctAnswer': answer,
            'explanation': str(raw.get('explanation') or '').strip(), 'difficulty': slot.difficulty,
            'topic': slot.topic, 'order': order, 'blueprintSlotId': str(slot.id),
            'questionType': question_type, 'score': float(slot.score), 'slotMetadata': slot.metadata or {}}


def generate_questions_with_ai(paper: ExamPaper, user_email: str, initial_rows: list[dict] | None = None, on_progress=None) -> list[dict]:
    version = paper.blueprint_version
    if not version or version.status != 'LOCKED':
        raise ValueError('Mỗi đề phải dùng một phiên bản ma trận đã khóa trước khi sinh AI.')
    slots = list(version.slots.all().order_by('position'))
    if not slots:
        raise ValueError('Phiên bản ma trận chưa có slot.')
    if paper.total_questions != len(slots):
        raise ValueError('Số câu của đề không khớp số slot trong phiên bản ma trận.')
    rows = list(initial_rows or [])
    if len(rows) > len(slots):
        rows = []
    answer_targets = _balanced_answer_targets(slots)
    for slot in slots[len(rows):]:
        previous_questions = [{'order': item['order'], 'content': item['content']} for item in rows[-12:]]
        prompt = f'''Đóng vai giáo viên Toán nhiều năm kinh nghiệm ra đề Olympic. Sinh DUY NHẤT một câu cho slot cố định sau; không đổi loại câu, số phương án, điểm, độ khó, chủ đề hay metadata bắt buộc.
Không sao chép và không tạo biến thể chỉ thay số. Thay đổi bối cảnh, dữ liệu biểu diễn và đường suy luận khi phù hợp CTX/Assessment Intent. Các câu đã sinh để tránh trùng: {json.dumps(previous_questions, ensure_ascii=False)}
Đề: {paper.title}; Cuộc thi: {paper.competition.name if paper.competition else ''}; Môn: {paper.subject}; Khối/Bảng: {paper.grade_or_category}; Ngôn ngữ: {paper.language}.
Slot: {json.dumps({'position': slot.position, 'questionType': slot.question_type, 'optionCount': slot.option_count, 'score': float(slot.score), 'difficulty': slot.difficulty, 'difficultyLabel': (slot.metadata or {}).get('difficultyLabel') or slot.difficulty, 'topic': slot.topic, 'knowledgeSource': slot.knowledge_source, 'knowledgeRequirements': slot.knowledge_requirements, 'prohibitedKnowledge': slot.prohibited_knowledge, 'assessmentIntent': slot.assessment_intent, 'estimatedSeconds': slot.estimated_seconds, 'metadata': slot.metadata}, ensure_ascii=False)}
Trong đó difficultyLabel là nhãn độ khó nguyên bản của file ma trận và là yêu cầu ưu tiên khi thiết kế độ phức tạp của câu; difficulty chỉ là mã nội bộ để kiểm tra hệ thống.
Trả JSON đúng dạng {{"question":{{"content":"","choices":[],"correctAnswer":"","explanation":""}}}}. Với numeric_input: choices phải là mảng rỗng và correctAnswer là đáp số. Không sao chép nguyên văn nguồn.
Nguồn tham chiếu:
{source_context(paper, 12000)}'''
        result = call_ai_json(paper=paper, task_type='generate_slot', model=_config().generation_model, user_email=user_email, system='Bạn là giáo viên ra đề kỳ cựu và biên tập viên khảo thí. Chỉ trả JSON hợp lệ; tuân thủ tuyệt đối slot ma trận.', prompt=prompt)
        row = result.get('question') if isinstance(result, dict) else None
        if not isinstance(row, dict):
            raise ValueError(f'AI không trả về câu hợp lệ cho slot {slot.position}.')
        candidate = normalize_question_for_slot(row, slot.position, slot)
        target = answer_targets.get(slot.position)
        if target:
            candidate = _place_correct_answer(candidate, target)
        duplicate_reasons = []
        for prior_position, prior in enumerate(rows, 1):
            duplicate, reason = _duplicate_pair(prior, candidate)
            if duplicate:
                duplicate_reasons.append(f'câu {prior_position}: {reason}')
        if duplicate_reasons:
            raise ValueError(f'Câu {slot.position} bị trùng hoặc chỉ thay số so với ' + '; '.join(duplicate_reasons[:3]) + '. Hãy sinh lại đề.')
        if len(set(_text_fingerprint(item) for item in candidate.get('choices', []))) != len(candidate.get('choices', [])):
            raise ValueError(f'Câu {slot.position} có phương án trùng nhau. Hãy sinh lại đề.')
        rows.append(candidate)
        if on_progress:
            on_progress(rows)
    return rows


def revise_question_with_ai(paper: ExamPaper, question: ExamQuestion, action: str, user_email: str) -> dict:
    slot = question.blueprint_slot
    if not slot:
        raise ValueError('Câu hỏi cũ chưa gắn slot ma trận; hãy tạo đề mới từ ma trận.')
    instructions = {
        'rewrite': 'Viết lại cho rõ ràng hơn nhưng giữ nguyên mọi ràng buộc của slot.',
        'distractors': 'Giữ nội dung và đáp án đúng, tạo phương án nhiễu hợp lý theo slot.',
        'increase-difficulty': 'Không được thay đổi độ khó vì độ khó đã khóa trong slot; cải thiện chất lượng diễn đạt.',
        'decrease-difficulty': 'Không được thay đổi độ khó vì độ khó đã khóa trong slot; cải thiện chất lượng diễn đạt.',
        'replace': 'Thay bằng câu khác nhưng giữ nguyên toàn bộ ràng buộc slot.',
    }
    if action not in instructions:
        raise ValueError('Thao tác AI cho câu hỏi không hợp lệ.')
    prompt = f'''{instructions[action]}
Slot bắt buộc: {json.dumps({'questionType': slot.question_type, 'optionCount': slot.option_count, 'score': float(slot.score), 'difficulty': slot.difficulty, 'topic': slot.topic, 'metadata': slot.metadata}, ensure_ascii=False)}
Trả JSON {{"question":{{"content":"","choices":[],"correctAnswer":"","explanation":""}}}}.
Câu hiện tại: {json.dumps(serialize_question(question), ensure_ascii=False)}
    Nguồn: {source_context(paper, 12000)}'''
    result = call_ai_json(paper=paper, task_type=f'question_{action}', model=_config().generation_model, user_email=user_email, system='Chỉ trả JSON hợp lệ và tuyệt đối không đổi ràng buộc của slot.', prompt=prompt)
    row = result.get('question') if isinstance(result, dict) else None
    if not isinstance(row, dict):
        raise ValueError('AI không trả về câu hỏi hợp lệ.')
    return normalize_question_for_slot(row, question.order, slot)


def revise_paper_from_chat(paper: ExamPaper, instruction: str, question_id: str, user_email: str) -> dict:
    """Turn a natural-language edit request into validated question changes."""
    message = str(instruction or '').strip()
    if not message:
        raise ValueError('Nhập yêu cầu cần chỉnh sửa.')
    if len(message) > 4000:
        raise ValueError('Yêu cầu chỉnh sửa tối đa 4.000 ký tự.')
    queryset = paper.questions.select_related('blueprint_slot').all()
    if question_id:
        queryset = queryset.filter(pk=question_id)
    questions = list(queryset)
    if not questions:
        raise ValueError('Không tìm thấy câu hỏi trong phạm vi chỉnh sửa.')
    current = [serialize_question(question) for question in questions]
    scope = f'Câu {questions[0].order}' if question_id else f'toàn bộ {len(questions)} câu của đề'
    prior_chat = [
        {'role': item.get('role'), 'text': str(item.get('text') or '')[:1000]}
        for item in (paper.ai_chat_history or [])[:-1][-10:]
        if isinstance(item, dict) and item.get('role') in {'user', 'assistant'} and item.get('text')
    ]
    prompt = f'''Người biên tập yêu cầu chỉnh sửa {scope}: {message}
Chỉ thay đổi những gì yêu cầu; giữ nguyên các phần không liên quan. Mọi câu phải tiếp tục tuân thủ slot ma trận, loại câu, số phương án, độ khó và điểm.
Trả JSON dạng {{"reply":"Mô tả ngắn thay đổi đã thực hiện","changes":[{{"id":"UUID hiện có","question":{{"content":"","choices":[],"correctAnswer":"","explanation":""}}}}]}}.
Nếu yêu cầu chỉ hỏi hoặc không cần sửa, trả changes là mảng rỗng và giải thích trong reply.
Trao đổi trước trong phiên chat của đề: {json.dumps(prior_chat, ensure_ascii=False)}
Câu hiện tại: {json.dumps(current, ensure_ascii=False)}
Nguồn tham chiếu rút gọn: {source_context(paper, 20000)}'''
    result = call_ai_json(
        paper=paper, task_type='paper_chat', model=_config().generation_model, user_email=user_email,
        system='Bạn là đồng biên tập đề thi. Chỉ trả JSON hợp lệ và không được phá vỡ ràng buộc ma trận.',
        prompt=prompt, max_tokens=3000 if question_id else 8000,
    )
    raw_changes = result.get('changes') if isinstance(result, dict) else []
    if not isinstance(raw_changes, list):
        raise ValueError('AI trả về danh sách chỉnh sửa không hợp lệ.')
    by_id = {str(question.id): question for question in questions}
    changes = []
    for raw_change in raw_changes:
        if not isinstance(raw_change, dict):
            continue
        target = by_id.get(str(raw_change.get('id') or ''))
        row = raw_change.get('question')
        if not target or not isinstance(row, dict):
            continue
        merged = {**serialize_question(target), **row, 'order': target.order}
        normalized = normalize_question_for_slot(merged, target.order, target.blueprint_slot) if target.blueprint_slot else normalize_question(merged, target.order)
        changes.append({'id': str(target.id), 'question': normalized})
    reply = str(result.get('reply') or '').strip() if isinstance(result, dict) else ''
    return {'reply': reply or ('Đã chuẩn bị các thay đổi.' if changes else 'Không có nội dung nào cần thay đổi.'), 'changes': changes}
