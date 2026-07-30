from __future__ import annotations

import hashlib
import json
import re
import unicodedata
from io import BytesIO

from django.db import transaction
from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response

from authentication.permissions import IsManagerOrAdmin
from .models import Blueprint, BlueprintVersion, Competition, ExamSession
from .paper_views import IsWorkspaceUser, actor
from .blueprint_services import (difficulty_distribution, draft_slots, lock_version, replace_slots,
    serialize_blueprint, serialize_version, validate_metadata_schema)


def blueprint_payload(data, current=None):
    source = data or {}
    competition_id = str(source.get('competitionId', getattr(current, 'competition_id', '') or '')).strip()
    session_id = str(source.get('sessionId', getattr(current, 'session_id', '') or '')).strip()
    competition = Competition.objects.filter(pk=competition_id).first() if competition_id else None
    session = ExamSession.objects.filter(pk=session_id).first() if session_id else None
    if competition_id and not competition:
        raise ValueError('Không tìm thấy cuộc thi đã chọn.')
    if session_id and not session:
        raise ValueError('Không tìm thấy kỳ thi đã chọn.')
    if session and not competition:
        competition = Competition.objects.filter(pk=session.competition_id).first()
    return {
        'name': str(source.get('name', getattr(current, 'name', '') or '')).strip(), 'competition': competition, 'session': session,
        'round_name': str(source.get('roundName', getattr(current, 'round_name', '') or '')).strip(),
        'subject': str(source.get('subject', getattr(current, 'subject', '') or '')).strip(),
        'grade_or_category': str(source.get('gradeOrCategory', getattr(current, 'grade_or_category', '') or '')).strip(),
        'language': str(source.get('language', getattr(current, 'language', 'Tiếng Việt') or 'Tiếng Việt')).strip(),
        'duration_minutes': max(1, int(source.get('durationMinutes', getattr(current, 'duration_minutes', 60)) or 60)),
        'metadata_schema': validate_metadata_schema(source.get('metadataSchema', getattr(current, 'metadata_schema', {}) or {})),
        'description': str(source.get('description', getattr(current, 'description', '') or '')).strip(),
    }


@api_view(['GET', 'POST'])
@permission_classes([IsWorkspaceUser])
def blueprints_list(request):
    if request.method == 'GET':
        rows = Blueprint.objects.select_related('competition', 'session').prefetch_related('versions__slots').all()
        competition_id = str(request.query_params.get('competitionId') or '').strip()
        if competition_id:
            rows = rows.filter(competition_id=competition_id)
        return Response({'items': [serialize_blueprint(item, include_versions=True) for item in rows[:500]]})
    if getattr(request, 'user_role', '') not in {'ADMIN', 'MANAGER'}:
        return Response({'error': 'Bạn không có quyền tạo ma trận đề.'}, status=status.HTTP_403_FORBIDDEN)
    try:
        payload = blueprint_payload(request.data)
        if not payload['name']:
            raise ValueError('Tên ma trận đề là bắt buộc.')
        with transaction.atomic():
            item = Blueprint.objects.create(**payload, created_by=actor(request), updated_by=actor(request))
            version = BlueprintVersion.objects.create(blueprint=item, version_number=1, note=str(request.data.get('versionNote') or ''), created_by=actor(request))
            if request.data.get('slots'):
                replace_slots(version, request.data.get('slots') or [])
        return Response(serialize_blueprint(Blueprint.objects.prefetch_related('versions__slots').get(pk=item.pk), include_versions=True), status=status.HTTP_201_CREATED)
    except (TypeError, ValueError) as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsWorkspaceUser])
def blueprint_detail(request, pk):
    item = Blueprint.objects.select_related('competition', 'session').prefetch_related('versions__slots').filter(pk=pk).first()
    if not item:
        return Response({'error': 'Không tìm thấy ma trận đề.'}, status=status.HTTP_404_NOT_FOUND)
    if request.method == 'GET':
        return Response(serialize_blueprint(item, include_versions=True))
    if getattr(request, 'user_role', '') not in {'ADMIN', 'MANAGER'}:
        return Response({'error': 'Bạn không có quyền chỉnh sửa ma trận đề.'}, status=status.HTTP_403_FORBIDDEN)
    if request.method == 'DELETE':
        if item.versions.filter(exam_papers__isnull=False).exists():
            return Response({'error': 'Không thể xóa ma trận đã được dùng để sinh đề.'}, status=status.HTTP_400_BAD_REQUEST)
        item.delete(); return Response({'success': True})
    try:
        payload = blueprint_payload(request.data, item)
        if not payload['name']:
            raise ValueError('Tên ma trận đề là bắt buộc.')
        for key, value in payload.items(): setattr(item, key, value)
        item.updated_by = actor(request); item.save()
        return Response(serialize_blueprint(Blueprint.objects.prefetch_related('versions__slots').get(pk=item.pk), include_versions=True))
    except (TypeError, ValueError) as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsManagerOrAdmin])
def blueprint_draft_from_config(request):
    try:
        total = int(request.data.get('totalQuestions') or 0)
        distribution = difficulty_distribution(total, request.data.get('difficultyDistribution'))
        topics = request.data.get('topics') if isinstance(request.data.get('topics'), list) else []
        slots = draft_slots(total, distribution, topics, str(request.data.get('questionType') or 'single_choice'), int(request.data.get('optionCount') or 4), str(request.data.get('description') or ''))
        payload = blueprint_payload({**request.data, 'name': request.data.get('name') or f'Ma trận nháp {request.data.get("subject") or ""}'.strip()})
        if not payload['name']:
            payload['name'] = 'Ma trận nháp'
        with transaction.atomic():
            item = Blueprint.objects.create(**payload, created_by=actor(request), updated_by=actor(request))
            version = BlueprintVersion.objects.create(blueprint=item, version_number=1, note='Tạo tự động từ wizard tạo đề', created_by=actor(request))
            replace_slots(version, slots)
        return Response({'blueprint': serialize_blueprint(Blueprint.objects.prefetch_related('versions__slots').get(pk=item.pk), include_versions=True), 'version': serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True)}, status=status.HTTP_201_CREATED)
    except (TypeError, ValueError) as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsManagerOrAdmin])
def blueprint_duplicate(request, pk):
    source = Blueprint.objects.prefetch_related('versions__slots').filter(pk=pk).first()
    if not source:
        return Response({'error': 'Không tìm thấy ma trận đề.'}, status=status.HTTP_404_NOT_FOUND)
    with transaction.atomic():
        copy = Blueprint.objects.create(name=f'{source.name} — Bản sao', competition=source.competition, session=source.session, round_name=source.round_name, subject=source.subject, grade_or_category=source.grade_or_category, language=source.language, metadata_schema=source.metadata_schema, description=source.description, created_by=actor(request), updated_by=actor(request))
        latest = source.versions.order_by('-version_number').first()
        version = BlueprintVersion.objects.create(blueprint=copy, version_number=1, note=f'Nhân bản từ {source.name}', created_by=actor(request))
        if latest:
            replace_slots(version, [{'position': slot.position, 'questionType': slot.question_type, 'optionCount': slot.option_count, 'score': str(slot.score), 'difficulty': slot.difficulty, 'topic': slot.topic, 'knowledgeSource': slot.knowledge_source, 'knowledgeRequirements': slot.knowledge_requirements, 'prohibitedKnowledge': slot.prohibited_knowledge, 'assessmentIntent': slot.assessment_intent, 'estimatedSeconds': slot.estimated_seconds, 'metadata': slot.metadata} for slot in latest.slots.all()])
    return Response(serialize_blueprint(Blueprint.objects.prefetch_related('versions__slots').get(pk=copy.pk), include_versions=True), status=status.HTTP_201_CREATED)


@api_view(['POST'])
@permission_classes([IsManagerOrAdmin])
def blueprint_new_version(request, pk):
    item = Blueprint.objects.prefetch_related('versions__slots').filter(pk=pk).first()
    if not item:
        return Response({'error': 'Không tìm thấy ma trận đề.'}, status=status.HTTP_404_NOT_FOUND)
    with transaction.atomic():
        newest = item.versions.order_by('-version_number').first()
        version = BlueprintVersion.objects.create(blueprint=item, version_number=(newest.version_number if newest else 0) + 1, note=str(request.data.get('note') or ''), created_by=actor(request))
        if newest:
            replace_slots(version, [{'position': slot.position, 'questionType': slot.question_type, 'optionCount': slot.option_count, 'score': str(slot.score), 'difficulty': slot.difficulty, 'topic': slot.topic, 'knowledgeSource': slot.knowledge_source, 'knowledgeRequirements': slot.knowledge_requirements, 'prohibitedKnowledge': slot.prohibited_knowledge, 'assessmentIntent': slot.assessment_intent, 'estimatedSeconds': slot.estimated_seconds, 'metadata': slot.metadata} for slot in newest.slots.all()])
    return Response(serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True), status=status.HTTP_201_CREATED)


@api_view(['GET', 'PUT'])
@permission_classes([IsWorkspaceUser])
def blueprint_version_detail(request, pk):
    version = BlueprintVersion.objects.select_related('blueprint__competition', 'blueprint__session').prefetch_related('slots').filter(pk=pk).first()
    if not version:
        return Response({'error': 'Không tìm thấy phiên bản ma trận.'}, status=status.HTTP_404_NOT_FOUND)
    if request.method == 'GET':
        return Response(serialize_version(version, include_slots=True))
    if getattr(request, 'user_role', '') not in {'ADMIN', 'MANAGER'}:
        return Response({'error': 'Bạn không có quyền chỉnh sửa phiên bản ma trận.'}, status=status.HTTP_403_FORBIDDEN)
    try:
        replace_slots(version, request.data.get('slots') or [])
        version.note = str(request.data.get('note', version.note) or '')
        version.save(update_fields=['note', 'updated_at'])
        return Response(serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True))
    except (TypeError, ValueError) as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsManagerOrAdmin])
def blueprint_version_lock(request, pk):
    version = BlueprintVersion.objects.prefetch_related('slots').filter(pk=pk).first()
    if not version:
        return Response({'error': 'Không tìm thấy phiên bản ma trận.'}, status=status.HTTP_404_NOT_FOUND)
    try:
        lock_version(version, actor(request))
        return Response(serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True))
    except ValueError as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)


def _header_map(sheet):
    fallback = ({}, 1)
    for row_number, values in enumerate(sheet.iter_rows(min_row=1, max_row=min(sheet.max_row, 30), values_only=True), 1):
        headers = {_plain(value): index for index, value in enumerate(values) if str(value or '').strip()}
        if row_number == 1:
            fallback = (headers, row_number)
        has_position = any(name in headers for name in ('vi tri', 'cau', 'stt', 'so thu tu', 'position', 'question'))
        has_difficulty = any(name in headers for name in ('muc do', 'do kho', 'difficulty', 'cap do', 'muc do nhan thuc'))
        if has_position and has_difficulty:
            return headers, row_number
    return fallback


def _value(row, headers, *names):
    for name in names:
        index = headers.get(_plain(name))
        if index is not None and index < len(row): return row[index]
    return ''


def _plain(value) -> str:
    text = unicodedata.normalize('NFD', str(value or '').strip().lower())
    text = ''.join(character for character in text if unicodedata.category(character) != 'Mn').replace('đ', 'd')
    return ' '.join(text.split())


def _difficulty_code(value, five_tier: bool = False) -> str:
    label = _plain(value)
    if label in {'easy', 'de', 'rat de', 'nhan biet', 'biet'}:
        return 'EASY'
    if label in {'medium', 'trung binh', 'thong hieu', 'hieu'}:
        return 'MEDIUM'
    if label in {'hard', 'kha', 'van dung'} or (label == 'kho' and not five_tier):
        return 'HARD'
    if label in {'very_hard', 'very hard', 'rat kho', 'van dung cao'} or (label == 'kho' and five_tier):
        return 'VERY_HARD'
    return 'MEDIUM'


def _question_type_code(value) -> str:
    label = _plain(value)
    if any(token in label for token in ('dien dap so', 'nhap dap so', 'tra loi ngan', 'numeric', 'short answer')):
        return 'numeric_input'
    return 'single_choice'


def _numeric(value, default=0):
    match = re.search(r'-?\d+(?:[.,]\d+)?', str(value or '').replace(' ', ''))
    return float(match.group().replace(',', '.')) if match else default


def _source_profile(uploaded) -> dict:
    content = uploaded.read()
    uploaded.seek(0)
    extension = str(uploaded.name or '').lower().rsplit('.', 1)[-1]
    return {
        'fileName': str(uploaded.name or ''), 'extension': extension,
        'sizeBytes': len(content), 'sha256': hashlib.sha256(content).hexdigest(),
        'importedAt': timezone.now().isoformat(),
    }


def _docx_blueprint_rows(uploaded, version) -> list[dict]:
    document = Document(BytesIO(uploaded.read()))
    allocation = {'single_choice': {'count': 0, 'points': 0.0}, 'numeric_input': {'count': 0, 'points': 0.0}}
    matrix_table = None
    header_row_index = 0
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            cells = [_plain(cell.text) for cell in row.cells]
            if {'phan', 'so cau', 'diem'}.issubset(set(cells)):
                for data_row in table.rows[row_index + 1:]:
                    values = [cell.text.strip() for cell in data_row.cells]
                    label = _plain(values[0] if values else '')
                    kind = 'numeric_input' if 'dien dap so' in label or 'nhap dap so' in label else 'single_choice' if 'chon dap an' in label else ''
                    if kind and len(values) >= 3:
                        allocation[kind] = {'count': int(float(values[1] or 0)), 'points': float(values[2] or 0)}
            has_position = any(cell in {'cau', 'stt', 'so thu tu', 'vi tri', 'position', 'question'} for cell in cells)
            has_difficulty = any(cell in {'muc do', 'do kho', 'difficulty', 'cap do', 'muc do nhan thuc'} for cell in cells)
            if has_position and has_difficulty:
                matrix_table, header_row_index = table, row_index
    if not matrix_table:
        raise ValueError('Không tìm thấy bảng câu hỏi có cột vị trí/STT và cột độ khó/mức độ.')
    headers = {_plain(cell.text): index for index, cell in enumerate(matrix_table.rows[header_row_index].cells)}
    def cell_value(values, *names):
        for name in names:
            index = headers.get(_plain(name))
            if index is not None and index < len(values):
                return values[index]
        return ''
    raw_rows = []
    for row in matrix_table.rows[header_row_index + 1:]:
        values = [cell.text.strip() for cell in row.cells]
        source_row = {matrix_table.rows[header_row_index].cells[index].text.strip(): value for index, value in enumerate(values) if index < len(matrix_table.rows[header_row_index].cells) and value}
        position_text = cell_value(values, 'Câu', 'STT', 'Số thứ tự', 'Vị trí', 'Position', 'Question')
        match = re.search(r'\d+', position_text)
        if not match:
            continue
        position = int(match.group())
        difficulty_label = cell_value(values, 'Mức độ', 'Độ khó', 'Difficulty', 'Cấp độ', 'Mức độ nhận thức') or 'Chưa phân loại'
        raw_rows.append({
            'position': position,
            'difficulty': _difficulty_code(difficulty_label),
            'topic': cell_value(values, 'Mạch kiến thức', 'Chủ đề', 'Topic', 'Nội dung', 'Đơn vị kiến thức'),
            'knowledgeSource': cell_value(values, 'Lớp', 'Khối', 'Nguồn kiến thức', 'Knowledge source'),
            'knowledgeRequirements': cell_value(values, 'Yêu cầu cần đạt', 'Yêu cầu kiến thức', 'Knowledge requirements', 'Mạch kiến thức', 'Nội dung'),
            'prohibitedKnowledge': cell_value(values, 'Không sử dụng', 'Kiến thức không sử dụng', 'Prohibited knowledge'),
            'assessmentIntent': cell_value(values, 'Assessment Intent', 'Mục tiêu đánh giá', 'Yêu cầu đánh giá'),
            'rawQuestionType': cell_value(values, 'Loại câu', 'Dạng câu', 'Hình thức', 'Question type'),
            'rawOptionCount': cell_value(values, 'Số phương án', 'PA', 'Option count'),
            'rawScore': cell_value(values, 'Điểm', 'Score', 'Số điểm'),
            'rawSeconds': cell_value(values, 'Thời gian dự kiến', 'Giây', 'Estimated seconds'),
            'metadata': {
                'difficultyLabel': difficulty_label,
                'questionStyle': cell_value(values, 'Kiểu câu'),
                'contextCode': cell_value(values, 'CTX', 'Ngữ cảnh', 'Context'),
                'CL': cell_value(values, 'CL'),
                'CaL': cell_value(values, 'CaL'),
                'VA': cell_value(values, 'VA'),
                'expectedMisconception': cell_value(values, 'Expected Misconception'),
                'sourceRow': source_row,
            },
        })
    raw_rows.sort(key=lambda item: item['position'])
    if not raw_rows:
        raise ValueError('Bảng ma trận không có dòng câu hỏi hợp lệ.')
    source_labels = {_plain(item['metadata'].get('difficultyLabel')) for item in raw_rows}
    five_tier = {'rat de', 'de', 'trung binh', 'kha', 'kho'}.issubset(source_labels)
    for item in raw_rows:
        item['difficulty'] = _difficulty_code(item['metadata'].get('difficultyLabel'), five_tier)
    numeric_count = allocation['numeric_input']['count']
    numeric_from = len(raw_rows) - numeric_count + 1 if numeric_count else len(raw_rows) + 1
    competition_code = _plain(getattr(version.blueprint.competition, 'code', ''))
    option_count = 5 if competition_code == 'fimo' else 4
    single_count = allocation['single_choice']['count'] or max(1, len(raw_rows) - numeric_count)
    single_score = allocation['single_choice']['points'] / single_count if allocation['single_choice']['points'] else 1
    numeric_score = allocation['numeric_input']['points'] / numeric_count if numeric_count and allocation['numeric_input']['points'] else 1
    for item in raw_rows:
        numeric = _question_type_code(item.pop('rawQuestionType')) == 'numeric_input' if item.get('rawQuestionType') else item['position'] >= numeric_from
        row_options = int(_numeric(item.pop('rawOptionCount'), option_count))
        row_score = _numeric(item.pop('rawScore'), numeric_score if numeric else single_score)
        row_seconds = int(_numeric(item.pop('rawSeconds'), 300 if numeric else 150))
        item.update({'questionType': 'numeric_input' if numeric else 'single_choice', 'optionCount': 0 if numeric else row_options, 'score': row_score, 'estimatedSeconds': row_seconds})
    return raw_rows


@api_view(['POST'])
@permission_classes([IsManagerOrAdmin])
def blueprint_import(request):
    uploaded = request.FILES.get('file')
    if not uploaded:
        return Response({'error': 'Chọn file ma trận XLSX hoặc DOCX.'}, status=status.HTTP_400_BAD_REQUEST)
    try:
        source = request.data.copy()
        filename = str(uploaded.name or '')
        if not source.get('name'):
            source['name'] = filename.rsplit('.', 1)[0]
        if not source.get('competitionId'):
            code = re.search(r'\b(FIMO|FIEO|FIAIO|FISO|IMO|IEO|ICO)\b', filename, re.IGNORECASE)
            competition = Competition.objects.filter(code__iexact=code.group(1)).first() if code else None
            if competition:
                source['competitionId'] = competition.pk
                session = ExamSession.objects.filter(competition_id=competition.pk).order_by('-time').first()
                if session:
                    source['sessionId'] = session.pk
                if competition.code.upper() == 'FIMO':
                    source.setdefault('subject', 'Toán')
                    source.setdefault('durationMinutes', 90)
        if not source.get('gradeOrCategory'):
            grade = re.search(r'(?:khối|khoi|lớp|lop)\s*(\d+)', filename, re.IGNORECASE)
            if grade:
                source['gradeOrCategory'] = f'Khối {grade.group(1)}'
        payload = blueprint_payload(source)
        if not payload['name']:
            raise ValueError('Tên ma trận đề là bắt buộc.')
        with transaction.atomic():
            item = Blueprint.objects.create(**payload, created_by=actor(request), updated_by=actor(request))
            version = BlueprintVersion.objects.create(blueprint=item, version_number=1, note=str(request.data.get('versionNote') or 'Nhập từ file ma trận'), created_by=actor(request))
            result = blueprint_version_import(request, version.pk)
            if result.status_code >= 400:
                raise ValueError(str(getattr(result, 'data', {}).get('error') or 'Không thể import ma trận.'))
        return Response({'blueprint': serialize_blueprint(Blueprint.objects.prefetch_related('versions__slots').get(pk=item.pk), include_versions=True), 'version': serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True)}, status=status.HTTP_201_CREATED)
    except (TypeError, ValueError) as exc:
        return Response({'error': str(exc)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsManagerOrAdmin])
def blueprint_version_import(request, pk):
    version = BlueprintVersion.objects.select_related('blueprint').filter(pk=pk).first()
    uploaded = request.FILES.get('file')
    if not version or not uploaded:
        return Response({'error': 'Cần chọn phiên bản nháp và file XLSX hoặc DOCX.'}, status=status.HTTP_400_BAD_REQUEST)
    if version.status != BlueprintVersion.STATUS_DRAFT:
        return Response({'error': 'Chỉ import vào phiên bản nháp.'}, status=status.HTTP_400_BAD_REQUEST)
    try:
        source_profile = _source_profile(uploaded)
        extension = str(uploaded.name).lower().rsplit('.', 1)[-1]
        if extension == 'docx':
            rows = _docx_blueprint_rows(uploaded, version)
            replace_slots(version, rows)
            version.refresh_from_db()
            version.analysis = {**(version.analysis or {}), 'source': source_profile}
            version.save(update_fields=['analysis', 'updated_at'])
            return Response(serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True))
        if extension != 'xlsx':
            raise ValueError('Chỉ hỗ trợ file XLSX hoặc DOCX.')
        workbook = load_workbook(BytesIO(uploaded.read()), read_only=True, data_only=True)
        sheet = workbook.active; headers, header_row = _header_map(sheet); rows = []
        for number, row in enumerate(sheet.iter_rows(min_row=header_row + 1, values_only=True), 1):
            if not any(value not in (None, '') for value in row): continue
            difficulty_label = _value(row, headers, 'Mức độ', 'Độ khó', 'Difficulty', 'Cấp độ', 'Mức độ nhận thức') or 'Chưa phân loại'
            raw_type = _value(row, headers, 'Loại câu', 'Dạng câu', 'Hình thức', 'Question type') or 'single_choice'
            question_type = _question_type_code(raw_type)
            source_row = {str(sheet.cell(header_row, index + 1).value or '').strip(): value for index, value in enumerate(row) if value not in (None, '')}
            rows.append({
                'position': _value(row, headers, 'Vị trí', 'Câu', 'STT', 'Position') or number,
                'questionType': question_type, 'optionCount': 0 if question_type == 'numeric_input' else int(_numeric(_value(row, headers, 'Số phương án', 'PA', 'Option count'), 4)),
                'score': _numeric(_value(row, headers, 'Điểm', 'Score', 'Số điểm'), 1),
                'difficulty': _difficulty_code(difficulty_label),
                'topic': _value(row, headers, 'Chủ đề', 'Mạch kiến thức', 'Nội dung', 'Đơn vị kiến thức', 'Topic'),
                'knowledgeSource': _value(row, headers, 'Nguồn kiến thức', 'Lớp', 'Khối', 'Knowledge source'),
                'knowledgeRequirements': _value(row, headers, 'Yêu cầu kiến thức', 'Yêu cầu cần đạt', 'Knowledge requirements'),
                'prohibitedKnowledge': _value(row, headers, 'Không sử dụng', 'Kiến thức không sử dụng', 'Prohibited knowledge'),
                'assessmentIntent': _value(row, headers, 'Assessment intent', 'Mục tiêu đánh giá', 'Yêu cầu đánh giá'),
                'estimatedSeconds': int(_numeric(_value(row, headers, 'Thời gian dự kiến', 'Giây', 'Estimated seconds'), 90)),
                'metadata': {'difficultyLabel': str(difficulty_label).strip(), 'sourceRow': source_row},
            })
        source_labels = {_plain(item['metadata'].get('difficultyLabel')) for item in rows}
        five_tier = {'rat de', 'de', 'trung binh', 'kha', 'kho'}.issubset(source_labels)
        for item in rows:
            item['difficulty'] = _difficulty_code(item['metadata'].get('difficultyLabel'), five_tier)
        replace_slots(version, rows)
        version.refresh_from_db()
        version.analysis = {**(version.analysis or {}), 'source': source_profile}
        version.save(update_fields=['analysis', 'updated_at'])
        return Response(serialize_version(BlueprintVersion.objects.prefetch_related('slots').get(pk=version.pk), include_slots=True))
    except Exception as exc:
        return Response({'error': f'Không thể import ma trận: {exc}'}, status=status.HTTP_400_BAD_REQUEST)


def _blueprint_xlsx(version: BlueprintVersion) -> bytes:
    workbook = Workbook(); sheet = workbook.active; sheet.title = 'Ma trận đề'
    profile = version.analysis or {}
    sheet.append(['Tên ma trận', version.blueprint.name])
    sheet.append(['Phiên bản', version.version_number])
    sheet.append(['Tổng số câu', profile.get('totalQuestions', 0)])
    sheet.append(['Tổng điểm', profile.get('totalScore', 0)])
    sheet.append(['Phân bố độ khó', json.dumps(profile.get('difficultyDistribution', {}), ensure_ascii=False)])
    sheet.append([])
    headers = ['Vị trí', 'Loại câu', 'Số phương án', 'Điểm', 'Độ khó', 'Chủ đề', 'Nguồn kiến thức', 'Yêu cầu kiến thức', 'Không sử dụng', 'Assessment Intent', 'Giây', 'Metadata JSON']
    sheet.append(headers)
    for cell in sheet[7]:
        cell.font = Font(bold=True, color='FFFFFF'); cell.fill = PatternFill('solid', fgColor='1F4FC9')
    for slot in version.slots.all():
        metadata = slot.metadata or {}
        sheet.append([slot.position, slot.question_type, slot.option_count, float(slot.score), metadata.get('difficultyLabel') or slot.difficulty, slot.topic, slot.knowledge_source, slot.knowledge_requirements, slot.prohibited_knowledge, slot.assessment_intent, slot.estimated_seconds, json.dumps(metadata, ensure_ascii=False)])
    widths = [10, 18, 14, 10, 18, 24, 22, 38, 30, 38, 10, 45]
    for index, width in enumerate(widths, 1): sheet.column_dimensions[chr(64 + index)].width = width
    sheet.freeze_panes = 'A8'; sheet.auto_filter.ref = f'A7:L{sheet.max_row}'
    output = BytesIO(); workbook.save(output); return output.getvalue()


def _blueprint_docx(version: BlueprintVersion) -> bytes:
    document = Document(); section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Mm(297), Mm(210)
    section.top_margin = section.bottom_margin = Mm(14)
    section.left_margin = section.right_margin = Mm(14)
    normal = document.styles['Normal']; normal.font.name = 'Times New Roman'; normal.font.size = Pt(9)
    title = document.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(version.blueprint.name.upper()); run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(16)
    profile = version.analysis or {}
    question_total = profile.get('totalQuestions', 0); score_total = profile.get('totalScore', 0)
    summary = document.add_paragraph(); summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(f'Phiên bản {version.version_number} · {question_total} câu · {score_total} điểm · {version.blueprint.duration_minutes} phút').bold = True
    distribution = ' · '.join(f'{label}: {count}' for label, count in (profile.get('difficultyDistribution') or {}).items())
    document.add_paragraph('Phân bố độ khó: ' + (distribution or 'Chưa phân loại'))
    headers = ['Câu', 'Loại', 'PA', 'Điểm', 'Độ khó', 'Chủ đề', 'Nguồn/Yêu cầu kiến thức', 'Không sử dụng', 'Assessment Intent', 'Giây']
    table = document.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]; cell.text = header
        for text_run in cell.paragraphs[0].runs: text_run.bold = True
    for slot in version.slots.all():
        metadata = slot.metadata or {}; cells = table.add_row().cells
        values = [slot.position, slot.question_type, slot.option_count, float(slot.score), metadata.get('difficultyLabel') or slot.difficulty, slot.topic, '\n'.join(filter(None, [slot.knowledge_source, slot.knowledge_requirements])), slot.prohibited_knowledge, slot.assessment_intent, slot.estimated_seconds]
        for index, value in enumerate(values):
            cells[index].text = str(value or '')
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for text_run in paragraph.runs: text_run.font.name = 'Times New Roman'; text_run.font.size = Pt(8)
    output = BytesIO(); document.save(output); return output.getvalue()


@api_view(['GET'])
@permission_classes([IsWorkspaceUser])
def blueprint_version_export(request, pk, export_type):
    version = BlueprintVersion.objects.select_related('blueprint').prefetch_related('slots').filter(pk=pk).first()
    if not version:
        return Response({'error': 'Không tìm thấy phiên bản ma trận.'}, status=status.HTTP_404_NOT_FOUND)
    name = slugify(version.blueprint.name) or 'ma-tran-de'
    if export_type == 'xlsx':
        response = HttpResponse(_blueprint_xlsx(version), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename={name}-v{version.version_number}.xlsx'; return response
    if export_type == 'docx':
        response = HttpResponse(_blueprint_docx(version), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={name}-v{version.version_number}.docx'; return response
    return Response({'error': 'Chỉ hỗ trợ xuất DOCX hoặc XLSX.'}, status=status.HTTP_400_BAD_REQUEST)
