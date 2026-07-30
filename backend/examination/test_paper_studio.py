from io import BytesIO
from types import SimpleNamespace
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from docx import Document
from openpyxl import load_workbook

from .blueprint_services import replace_slots
from .blueprint_views import _blueprint_docx, _blueprint_xlsx, _docx_blueprint_rows, _header_map
from .models import AiProviderConfig, Blueprint, BlueprintSlot, BlueprintVersion, Competition, ExamPaper, ExamQuestion
from .paper_services import _balanced_answer_targets, call_ai_json, document_export, generate_questions_with_ai, paper_quality_report, read_uploaded_source, revise_paper_from_chat


def docx_upload(document: Document, name: str = 'source.docx') -> SimpleUploadedFile:
    output = BytesIO()
    document.save(output)
    return SimpleUploadedFile(name, output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


class BlueprintDocxImportTests(TestCase):
    def test_fimo_matrix_docx_becomes_variable_question_slots(self):
        document = Document()
        allocation = document.add_table(rows=4, cols=3)
        allocation.rows[0].cells[0].text, allocation.rows[0].cells[1].text, allocation.rows[0].cells[2].text = 'Phần', 'Số câu', 'Điểm'
        allocation.rows[1].cells[0].text, allocation.rows[1].cells[1].text, allocation.rows[1].cells[2].text = 'Trắc nghiệm chọn đáp án', '2', '6'
        allocation.rows[2].cells[0].text, allocation.rows[2].cells[1].text, allocation.rows[2].cells[2].text = 'Trắc nghiệm điền đáp số', '1', '5'
        allocation.rows[3].cells[0].text, allocation.rows[3].cells[1].text, allocation.rows[3].cells[2].text = 'Tổng', '3', '11'
        matrix = document.add_table(rows=4, cols=10)
        headers = ['Câu', 'Lớp', 'Mạch kiến thức', 'Mức độ', 'Kiểu câu', 'Assessment Intent', 'CL', 'CaL', 'VA', 'Expected Misconception']
        for index, value in enumerate(headers): matrix.rows[0].cells[index].text = value
        values = [
            ['1', '6', 'Số tự nhiên', 'Rất dễ', 'Thuần toán', 'Tính toán', '1', '1', 'K', 'Sai phép tính'],
            ['2', '7', 'Hình học', 'Khá', 'Thực tế', 'Mô hình hóa', '3', '2', 'M', 'Nhầm đơn vị'],
            ['3', '6+7', 'Tổng hợp', 'Khó', 'Điền đáp số', 'Lập luận', '5', '2', 'L', 'Sai bước cuối'],
        ]
        for row_index, row in enumerate(values, 1):
            for column_index, value in enumerate(row): matrix.rows[row_index].cells[column_index].text = value
        version = SimpleNamespace(blueprint=SimpleNamespace(competition=SimpleNamespace(code='FIMO')))

        rows = _docx_blueprint_rows(docx_upload(document, 'matrix.docx'), version)

        self.assertEqual(len(rows), 3)
        self.assertEqual(rows[0]['optionCount'], 5)
        self.assertEqual(rows[0]['score'], 3)
        self.assertEqual(rows[1]['difficulty'], 'HARD')
        self.assertEqual(rows[2]['questionType'], 'numeric_input')
        self.assertEqual(rows[2]['score'], 5)
        self.assertEqual(rows[2]['metadata']['VA'], 'L')
        self.assertEqual(rows[0]['metadata']['difficultyLabel'], 'Rất dễ')
        self.assertEqual(rows[1]['metadata']['sourceRow']['Assessment Intent'], 'Mô hình hóa')

    def test_generic_matrix_preserves_original_difficulty_labels(self):
        document = Document()
        matrix = document.add_table(rows=5, cols=7)
        headers = ['STT', 'Độ khó', 'Nội dung', 'Hình thức', 'Số phương án', 'Điểm', 'Mục tiêu đánh giá']
        for index, value in enumerate(headers):
            matrix.rows[0].cells[index].text = value
        values = [
            ['1', 'Nhận biết', 'Phân số', 'Trắc nghiệm', '4', '1', 'Nhận biết khái niệm'],
            ['2', 'Vận dụng cao', 'Hình học', 'Trắc nghiệm', '4', '2', 'Lập luận nhiều bước'],
            ['3', 'Mức A riêng', 'Dữ liệu', 'Trả lời ngắn', '0', '3', 'Giải quyết vấn đề'],
            ['4', 'Khó', 'Tổ hợp', 'Trắc nghiệm', '4', '2', 'Vận dụng nhiều bước'],
        ]
        for row_index, row in enumerate(values, 1):
            for column_index, value in enumerate(row):
                matrix.rows[row_index].cells[column_index].text = value
        version = SimpleNamespace(blueprint=SimpleNamespace(competition=SimpleNamespace(code='OTHER')))

        rows = _docx_blueprint_rows(docx_upload(document, 'generic.docx'), version)

        self.assertEqual([row['metadata']['difficultyLabel'] for row in rows], ['Nhận biết', 'Vận dụng cao', 'Mức A riêng', 'Khó'])
        self.assertEqual([row['difficulty'] for row in rows], ['EASY', 'VERY_HARD', 'MEDIUM', 'HARD'])
        self.assertEqual(rows[2]['questionType'], 'numeric_input')
        self.assertEqual(rows[1]['score'], 2)

    def test_docx_source_reader_includes_table_cells(self):
        document = Document()
        document.add_paragraph('Yêu cầu chung')
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text, table.cell(0, 1).text = 'Câu', 'Mức độ'
        table.cell(1, 0).text, table.cell(1, 1).text = '1', 'Khó'

        extracted = read_uploaded_source(docx_upload(document))

        self.assertIn('Yêu cầu chung', extracted)
        self.assertIn('Câu | Mức độ', extracted)
        self.assertIn('1 | Khó', extracted)


class BlueprintProfileExportTests(TestCase):
    def setUp(self):
        competition = Competition.objects.create(id='profile-test', code='TEST', name='Test', parent='FT', organizer='FT', sort_key='profile-test')
        blueprint = Blueprint.objects.create(name='Ma trận linh hoạt', competition=competition, subject='Toán', grade_or_category='Khối 8', duration_minutes=75)
        self.version = BlueprintVersion.objects.create(blueprint=blueprint, version_number=1)
        replace_slots(self.version, [
            {'position': 1, 'questionType': 'single_choice', 'optionCount': 4, 'score': 1, 'difficulty': 'EASY', 'difficultyLabel': 'Nhận biết', 'topic': 'Số học', 'assessmentIntent': 'Kiểm tra khái niệm'},
            {'position': 2, 'questionType': 'single_choice', 'optionCount': 4, 'score': 2, 'difficulty': 'HARD', 'difficultyLabel': 'Vận dụng', 'topic': 'Đại số', 'assessmentIntent': 'Vận dụng'},
            {'position': 3, 'questionType': 'numeric_input', 'optionCount': 0, 'score': 3, 'difficulty': 'VERY_HARD', 'difficultyLabel': 'Vận dụng cao', 'topic': 'Hình học', 'assessmentIntent': 'Lập luận'},
        ])
        self.version.refresh_from_db()

    def test_profile_is_saved_once_with_exact_source_labels(self):
        self.assertEqual(self.version.analysis['totalQuestions'], 3)
        self.assertEqual(self.version.analysis['totalScore'], 6.0)
        self.assertEqual(self.version.analysis['difficultyDistribution'], {'Nhận biết': 1, 'Vận dụng': 1, 'Vận dụng cao': 1})
        self.assertEqual(self.version.analysis['questionTypeDistribution'], {'single_choice': 2, 'numeric_input': 1})

    def test_profile_exports_to_word_and_excel(self):
        word = Document(BytesIO(_blueprint_docx(self.version)))
        word_text = '\n'.join(paragraph.text for paragraph in word.paragraphs)
        self.assertIn('Nhận biết: 1', word_text)
        self.assertIn('Vận dụng cao: 1', word_text)

        workbook = load_workbook(BytesIO(_blueprint_xlsx(self.version)), data_only=True)
        sheet = workbook.active
        headers, header_row = _header_map(sheet)
        self.assertEqual(sheet['B3'].value, 3)
        self.assertIn('"Nhận biết": 1', sheet['B5'].value)
        self.assertEqual(sheet['E8'].value, 'Nhận biết')
        self.assertEqual(header_row, 7)
        self.assertEqual(headers['do kho'], 4)


class PaperPreviewExportTests(TestCase):
    def setUp(self):
        competition = Competition.objects.create(id='fimo-studio', code='FIMO', name='FIMO', parent='FT', organizer='FT', sort_key='fimo-studio')
        blueprint = Blueprint.objects.create(name='FIMO 7', competition=competition, round_name='Vòng loại Quốc gia', subject='Toán', grade_or_category='Khối 7')
        version = BlueprintVersion.objects.create(blueprint=blueprint, version_number=1, status='LOCKED')
        self.choice_slot = BlueprintSlot.objects.create(version=version, position=1, question_type='single_choice', option_count=5, score=3, difficulty='EASY', topic='Số học')
        self.numeric_slot = BlueprintSlot.objects.create(version=version, position=2, question_type='numeric_input', option_count=0, score=5, difficulty='VERY_HARD', topic='Tổng hợp')
        self.paper = ExamPaper.objects.create(title='Đề thử FIMO 7', competition=competition, blueprint_version=version, subject='Toán', grade_or_category='Khối 7', duration_minutes=90, total_questions=2)
        self.choice = ExamQuestion.objects.create(paper=self.paper, blueprint_slot=self.choice_slot, order=1, question_type='single_choice', score=3, content='Tính 12 + 8.', choices=['18', '19', '20', '21', '22'], correct_answer='C', explanation='12 + 8 = 20.', difficulty='EASY', topic='Số học')
        self.numeric = ExamQuestion.objects.create(paper=self.paper, blueprint_slot=self.numeric_slot, order=2, question_type='numeric_input', score=5, content='Tìm số còn thiếu.', choices=[], correct_answer='42', explanation='Suy luận được 42.', difficulty='VERY_HARD', topic='Tổng hợp')

    def test_word_export_matches_preview_structure_without_ai(self):
        payload = document_export(self.paper, 'combined')
        exported = Document(BytesIO(payload))
        text = '\n'.join(paragraph.text for paragraph in exported.paragraphs)

        self.assertGreater(len(payload), 10000)
        self.assertIn('ĐỀ THI VÒNG LOẠI QUỐC GIA', text)
        self.assertIn('E. 22', text)
        self.assertIn('Đáp số:', text)
        self.assertIn('HƯỚNG DẪN GIẢI', text)
        self.assertAlmostEqual(exported.sections[0].page_width.mm, 210, delta=0.2)
        self.assertAlmostEqual(exported.sections[0].page_height.mm, 297, delta=0.2)

    def test_draft_export_is_watermarked_but_official_export_is_not(self):
        draft = Document(BytesIO(document_export(self.paper, 'paper', 'draft')))
        official = Document(BytesIO(document_export(self.paper, 'paper', 'official')))
        draft_text = '\n'.join(paragraph.text for paragraph in draft.paragraphs)
        official_text = '\n'.join(paragraph.text for paragraph in official.paragraphs)
        self.assertIn('BẢN NHÁP — KHÔNG PHÁT HÀNH', draft_text)
        self.assertNotIn('BẢN NHÁP — KHÔNG PHÁT HÀNH', official_text)

    def test_answer_targets_are_evenly_distributed(self):
        slots = [SimpleNamespace(position=index + 1, question_type='single_choice', option_count=5) for index in range(30)]
        answers = list(_balanced_answer_targets(slots).values())
        self.assertEqual({label: answers.count(label) for label in 'ABCDE'}, {label: 6 for label in 'ABCDE'})
        self.assertTrue(all(answers[index] != answers[index - 1] for index in range(1, len(answers))))

    def test_quality_report_detects_number_only_variants(self):
        rows = [
            {'content':'Một cửa hàng có 12 hộp, mỗi hộp 5 bút. Có tất cả bao nhiêu bút?', 'choices':['50','55','60','65','70'], 'correctAnswer':'C', 'questionType':'single_choice', 'difficulty':'EASY'},
            {'content':'Một cửa hàng có 18 hộp, mỗi hộp 7 bút. Có tất cả bao nhiêu bút?', 'choices':['112','119','126','133','140'], 'correctAnswer':'C', 'questionType':'single_choice', 'difficulty':'VERY_HARD'},
        ]
        report = paper_quality_report(self.paper, rows)
        self.assertTrue(report['duplicatePairs'])
        self.assertEqual(report['duplicatePairs'][0]['reason'], 'chỉ thay số')

    @patch('examination.paper_services.call_ai_json')
    def test_blueprint_generation_places_answer_and_keeps_distinct_questions(self, mock_ai):
        mock_ai.side_effect = [
            {'question': {'content':'Tính tổng của 12 và 8.', 'choices':['18','19','20','21','22'], 'correctAnswer':'C', 'explanation':'Kết quả bằng 20.'}},
            {'question': {'content':'Tìm số tự nhiên thỏa mãn điều kiện đã cho.', 'choices':[], 'correctAnswer':'42', 'explanation':'Suy luận được 42.'}},
        ]
        rows = generate_questions_with_ai(self.paper, 'editor@example.com')
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['correctAnswer'], 'A')
        self.assertEqual(rows[0]['choices'][0], '20')
        self.assertEqual(rows[1]['correctAnswer'], '42')

    @patch('examination.paper_services.call_ai_json')
    def test_blueprint_generation_resumes_after_saved_question(self, mock_ai):
        mock_ai.return_value = {'question': {'content':'Tìm số tự nhiên thỏa mãn điều kiện đã cho.', 'choices':[], 'correctAnswer':'42', 'explanation':'Suy luận được 42.'}}
        saved = [{
            'order': 1, 'blueprintSlotId': str(self.choice_slot.id), 'content': 'Tính tổng của 12 và 8.',
            'choices': ['20', '18', '19', '21', '22'], 'correctAnswer': 'A', 'explanation': 'Kết quả bằng 20.',
            'difficulty': 'EASY', 'topic': 'Số học', 'questionType': 'single_choice', 'score': 3,
            'slotMetadata': {},
        }]
        progress = []

        rows = generate_questions_with_ai(self.paper, 'editor@example.com', initial_rows=saved, on_progress=lambda value: progress.append(len(value)))

        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['content'], saved[0]['content'])
        self.assertEqual(progress, [2])
        mock_ai.assert_called_once()

    @patch('examination.paper_services.call_ai_json')
    def test_chat_revision_is_normalized_against_the_locked_slot(self, mock_ai):
        mock_ai.return_value = {'reply': 'Đã làm rõ câu dẫn.', 'changes': [{'id': str(self.choice.id), 'question': {'content': 'Tính tổng 12 và 8.', 'choices': ['18', '19', '20', '21', '22'], 'correctAnswer': 'C', 'explanation': '12 + 8 = 20.'}}]}

        result = revise_paper_from_chat(self.paper, 'Làm rõ câu dẫn', str(self.choice.id), 'editor@example.com')

        self.assertEqual(result['reply'], 'Đã làm rõ câu dẫn.')
        self.assertEqual(result['changes'][0]['question']['content'], 'Tính tổng 12 và 8.')
        self.assertEqual(len(result['changes'][0]['question']['choices']), 5)
        self.assertEqual(result['changes'][0]['question']['difficulty'], 'EASY')


class AiConfigFailoverTests(TestCase):
    def setUp(self):
        competition = Competition.objects.create(id='ai-failover', code='AI', name='AI', parent='FT', organizer='FT', sort_key='ai-failover')
        self.paper = ExamPaper.objects.create(title='Đề thử AI', competition=competition, subject='Toán', grade_or_category='Khối 7', total_questions=1)
        self.primary = AiProviderConfig.objects.create(name='Key chính', priority=1, api_key_encrypted='primary-key')
        self.backup = AiProviderConfig.objects.create(name='Key dự phòng', priority=2, api_key_encrypted='backup-key')

    @patch('examination.paper_services.requests.post')
    @patch('examination.paper_services.decrypt_secret')
    def test_quota_exhaustion_marks_primary_and_uses_backup(self, mock_decrypt, mock_post):
        mock_decrypt.side_effect = ['primary-key', 'backup-key']
        exhausted = SimpleNamespace(
            ok=False, status_code=429, text='quota exhausted',
            json=lambda: {'error': {'code': 'credit_balance_exhausted', 'type': 'insufficient_quota', 'message': 'No credits'}},
        )
        success = SimpleNamespace(
            ok=True, status_code=200,
            json=lambda: {'choices': [{'message': {'content': '{"ok": true}'}}], 'usage': {'prompt_tokens': 10, 'completion_tokens': 2}},
        )
        mock_post.side_effect = [exhausted, success]

        result = call_ai_json(paper=self.paper, task_type='generate', model='fallback-model', system='system', prompt='prompt', user_email='editor@example.com')

        self.primary.refresh_from_db()
        self.backup.refresh_from_db()
        self.assertEqual(result, {'ok': True})
        self.assertEqual(self.primary.health_status, AiProviderConfig.STATUS_EXHAUSTED)
        self.assertEqual(self.backup.health_status, AiProviderConfig.STATUS_READY)
        self.assertIn('Key dự phòng', self.paper._ai_route_notice)
        self.assertEqual(mock_post.call_count, 2)
